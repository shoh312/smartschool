# -*- coding: utf-8 -*-
"""Writes the two expert reviews (тақриз) for SmartFlow.

These are the documents an expert signs, so the text must describe this
system and no other. The earlier copies still spoke of Django and document
management -- the previous applicant's product -- and would have been read
by the signatory as a description of something that does not exist here.

Two reviews, written for two different readers on purpose: the first from a
technical standpoint, the second from the standpoint of what it changes for
a school. That is how the template pair was split, and it is the right
split -- a second technical review adds nothing.

The signature block is deliberately left blank. Nobody may be named as a
reviewer until they have read the system and agreed to sign.

    python docs/registration/build_reviews.py
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tayyor")

# The author is named; the reviewer is not. A reviewer's name may only be
# printed once that person has read the system and agreed to sign it.
def _author():
    """Same source as fill_documents.py -- see the note there."""
    import json
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "applicant.json")
    if not os.path.exists(path):
        raise SystemExit("applicant.json topilmadi.")
    with open(path, encoding="utf-8") as handle:
        return json.load(handle)["fish"]


AUTHOR = _author()


def new_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
    return doc


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)


def para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def signature(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Тақриздиҳанда:").bold = True
    for line in ("_______________________________________________",
                 "(унвони илмӣ, вазифа, муассиса)",
                 "",
                 "_______________________  /  _______________________",
                 "(имзо)                                  (насаб ва ҳарфҳои аввал)",
                 "",
                 "«____» ______________ 2026 с."):
        q = doc.add_paragraph(line)
        q.paragraph_format.space_after = Pt(2)


# ------------------------------------------------------------- review 1 --

def technical_review():
    doc = new_document()
    title(doc, "Тақриз ба системаи «SmartFlow»")

    p = doc.add_paragraph()
    p.add_run("Таҳиягар: ").bold = True
    p.add_run(AUTHOR)

    para(doc,
         "«SmartFlow» — системаи худкори бақайдгирии ҳозиршавии хонандагон "
         "ва идоракунии раванди таълим мебошад. Ҳозиршавӣ ба таври худкор "
         "муайян карда мешавад: камераҳои IP, ки дар синфхонаҳо насб "
         "шудаанд, тасвирро ба сервери мактабӣ мефиристанд, ки дар он ҷо "
         "барнома чеҳраи хонандагонро шинохта, ҳузур, дер омадан ё "
         "набудани онҳоро мустақилона қайд мекунад. Иштироки муаллим дар "
         "ин раванд лозим нест.")

    para(doc,
         "Система аз се қисм иборат аст: барномаи мобилӣ дар асоси Flutter, "
         "сервери мактабӣ ва сервери кушод, ки ҳарду дар асоси Python ва "
         "FastAPI сохта шудаанд, бо базаи додаҳои PostgreSQL. Шиносоии "
         "чеҳра тавассути модели InsightFace (buffalo_l) ва ONNX Runtime "
         "иҷро мешавад.")

    p = doc.add_paragraph()
    p.add_run("Аҳамияти илмӣ ва техникӣ.").bold = True

    para(doc,
         "Аз ҷиҳати техникӣ, кори зерин диққати махсусро сазовор аст. "
         "Танзимоти маъмулии моделҳои шиносоии чеҳра барои чеҳраи наздик ба "
         "объектив пешбинӣ шудаанд. Дар ин кор андозаи детектор то 960×960 "
         "нуқта афзоиш дода шуда, коэффитсиенти хурдкунии кадр тавре интихоб "
         "шудааст, ки чеҳраи хонандагони қаторҳои дур миқдори кофии пикселро "
         "нигоҳ дорад. Аломати биометрӣ вектори 512-андоза буда, муқоиса бо "
         "ченаки косинусии монандӣ анҷом дода мешавад.")

    para(doc,
         "Барои пешгирии хатоҳои шиносоӣ, ғайр аз ҳадди монандӣ, шарти "
         "иловагӣ ҷорӣ шудааст: варианти ғолиб бояд аз варианти дуюм бо "
         "фарқи муайян пеш гузарад. Ин кадрҳоеро, ки дар онҳо ду хонанда "
         "якхела эҳтимол доранд, рад мекунад — ба ҷои он ки формалӣ "
         "беҳтаринашро интихоб кунад. Илова бар ин, муқоиса танҳо бо "
         "рӯйхати хонандагони ҳамон синф иҷро мешавад, ки ба он камера "
         "нигаронида шудааст, ва ин дақиқиро новобаста аз шумораи умумии "
         "хонандагони мактаб нигоҳ медорад.")

    para(doc,
         "Қайди «ғоиб» танҳо пас аз ду давраи пайдарпайи шиносоӣ гузошта "
         "мешавад. Як давраи даҳсонияавӣ далели набудан нест: хонанда "
         "метавонад ҳамон лаҳза ворид шавад ё рӯяшро аз объектив гардонад. "
         "Ин ҳалли методӣ хатогии қайди нодурустро, ки боиси огоҳии бардурӯғ "
         "ба волидайн мегардад, бартараф мекунад.")

    para(doc,
         "Реҷаи сарфакоронаи кори камера низ қобили таваҷҷуҳ аст: камера "
         "танҳо дар давоми давраи шиносоӣ пайваст шуда, дар вақти дигар "
         "хомӯш мешавад. Ин зиёда аз навад фоизи вақти протсессорро озод "
         "мегузорад ва имкон медиҳад система дар компютери оддӣ бидуни "
         "корти графикӣ кор кунад.")

    p = doc.add_paragraph()
    p.add_run("Ҳифзи маълумоти шахсӣ.").bold = True

    para(doc,
         "Хусусияти принсипиалии меъмории система тақсимоти маълумот "
         "мебошад. Суратҳои хонандагон, аломатҳои биометрии онҳо ва "
         "тасвири камераҳо танҳо дар сервери мактабӣ, ки дар шабакаи "
         "дохилии муассиса ҷойгир аст, нигоҳ дошта ва коркард мешаванд ва "
         "ба шабакаи Интернет фиристода намешаванд. Ба берун танҳо натиҷа — "
         "қайди ҳозиршавӣ, баҳо ё эълон — мегузарад. Ин ҳал эътирози "
         "асосиро нисбати ҷорӣ намудани системаҳои биометрӣ дар мактабҳо "
         "бартараф мекунад.")

    p = doc.add_paragraph()
    p.add_run("Имкониятҳои иловагӣ.").bold = True

    para(doc,
         "Ғайр аз бақайдгирии ҳозиршавӣ, система журнали электронии баҳоҳо, "
         "рӯзнома, ҷадвали дарсҳо, супоришҳои онлайн ва таҳлили "
         "азхудкуниро дар бар мегирад. Зеҳни сунъӣ барои хондани баҳоҳо аз "
         "сурати журнали коғазӣ ва барои сохтани маводи таълимӣ аз рӯи "
         "мавзӯи додашуда истифода мешавад.")

    p = doc.add_paragraph()
    p.add_run("Хулоса.").bold = True

    para(doc,
         "«SmartFlow» кори мукаммали техникӣ буда, ҳаҷми он 40 982 сатри "
         "коди барномавӣ, 45 экрани барномаи мобилӣ ва 106 нуқтаи "
         "барномавии серверро ташкил медиҳад. Система дар ҳаҷми пурра "
         "татбиқ шуда, ба истифодаи озмоишӣ дар муассисаи таълимӣ омода "
         "аст. Кор ба бақайдгирии давлатӣ ҳамчун маҳсули зеҳнӣ сазовор "
         "мебошад.")

    signature(doc)
    return doc


# ------------------------------------------------------------- review 2 --

def practical_review():
    doc = new_document()
    title(doc, "Тақриз ба системаи «SmartFlow»")

    p = doc.add_paragraph()
    p.add_run("Таҳиягар: ").bold = True
    p.add_run(AUTHOR)

    para(doc,
         "«SmartFlow» — системаест, ки бақайдгирии ҳозиршавии хонандагон, "
         "журнали электронии баҳоҳо ва иттилоъдиҳии волидайнро дар як "
         "барномаи мобилӣ муттаҳид мекунад. Ҳозиршавӣ тавассути камера ва "
         "шиносоии чеҳра ба таври худкор қайд карда мешавад.")

    p = doc.add_paragraph()
    p.add_run("Мушкилоте, ки система ҳал мекунад.").bold = True

    para(doc,
         "Дар аксари муассисаҳои таълимӣ ҳозиршавӣ дастӣ гирифта мешавад. "
         "Ин вақти ҳар дарсро мегирад, манбаи хатогиҳо боқӣ мемонад ва "
         "маълумоти ҷамъшуда барои таҳлил истифода намешавад. Волидайн дар "
         "бораи набудани фарзандашон бо таъхир ё умуман хабар намеёбанд. "
         "Маъмурият намебинад, ки кадом синф ё кадом фан ақиб мемонад.")

    para(doc,
         "Ҷорӣ намудани системаҳои биометрӣ дар мактабҳо бо ду сабаб "
         "маҳдуд аст: арзиши баланди таҷҳизоти махсус ва нигаронии асосноки "
         "волидайн нисбати интиқоли суратҳои фарзандонашон ба системаҳои "
         "беруна. Системаи мазкур ҳарду монеаро бартараф мекунад.")

    p = doc.add_paragraph()
    p.add_run("Афзалиятҳои амалӣ.").bold = True

    for line in (
        "Барои кор камераи оддии IP кифоя аст — турникет ва таҷҳизоти "
        "махсус лозим нест.",
        "Шиносоӣ дар протсессори марказӣ иҷро мешавад — харидани корти "
        "графикӣ талаб карда намешавад.",
        "Суратҳо ва тасвири камера мактабро тарк намекунанд; ба Интернет "
        "танҳо натиҷа мегузарад.",
        "Барнома дар се забон — тоҷикӣ, русӣ ва англисӣ — кор мекунад.",
        "Волидайн бо нишон додани вақти дақиқ огоҳинома мегиранд: "
        "фарзанд соати чанд ба мактаб омад ё дер кард.",
        "Ҳангоми набудани Интернет сервери мактабӣ бақайдгирии ҳозиршавиро "
        "давом медиҳад.",
    ):
        q = doc.add_paragraph(line, style="List Bullet")
        q.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.add_run("Самаранокии иқтисодӣ.").bold = True

    para(doc,
         "Талаботи техникии система паст аст: компютери оддӣ бо протсессори "
         "чорядрагӣ ва хотираи 8 ГБ кифоя мебошад. Ин арзиши ҷорӣ намуданро "
         "барои муассисаҳои таълимӣ дастрас мегардонад. Система дар як "
         "вақт бо якчанд мактаб кор карда метавонад ва маълумоти ҳар "
         "мактаб пурра ҷудо нигоҳ дошта мешавад, ки ин имкони пешниҳоди "
         "хизматрасониро ба гурӯҳи мактабҳо фароҳам меорад.")

    p = doc.add_paragraph()
    p.add_run("Аудиторияи истифодабарандагон.").bold = True

    para(doc,
         "Система барои чор гурӯҳи корбарон пешбинӣ шудааст: директор ва "
         "маъмурияти мактаб, омӯзгорон, волидайн ва хонандагон. Ҳар кадом "
         "танҳо маълумоти ба ӯ дахлдорро мебинад.")

    p = doc.add_paragraph()
    p.add_run("Хулоса.").bold = True

    para(doc,
         "«SmartFlow» системаи инноватсионӣ буда, раванди маъмурияти "
         "мактабро содда ва самаранок мегардонад, вақти омӯзгоронро сарфа "
         "мекунад, ба волидайн маълумоти саривақтӣ медиҳад ва ба "
         "маъмурият асоси воқеӣ барои қабули қарор фароҳам меорад. "
         "Ҳамзамон, меъмории он маълумоти шахсии хонандагонро дар доираи "
         "муассиса нигоҳ медорад. Кор ба бақайдгирии давлатӣ ҳамчун "
         "маҳсули зеҳнӣ тавсия дода мешавад.")

    signature(doc)
    return doc


if __name__ == "__main__":
    os.makedirs(OUTPUT, exist_ok=True)
    for name, builder in (("SmartFlow_тақриз_1_техникӣ.docx", technical_review),
                          ("SmartFlow_тақриз_2_амалӣ.docx", practical_review)):
        path = os.path.join(OUTPUT, name)
        builder().save(path)
        print("Yozildi: %s" % name)
