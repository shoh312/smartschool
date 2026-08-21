# -*- coding: utf-8 -*-
"""Fills the ММПИ registration papers for SmartFlow.

The five documents another applicant used are the working templates: the
forms are the centre's own, so their layout and field numbering have to be
preserved exactly. Rather than retyping them, each file is copied and only
the values are replaced -- the tables, fonts and numbering stay untouched.

Personal details live in APPLICANT below and nowhere else. Everything about
the product is measured from the repository, not estimated.

    python docs/registration/fill_documents.py
"""

import os
import shutil

import docx

# ----------------------------------------------------------------- input --

TEMPLATES = r"C:\Users\gameboy\Downloads\Telegram Desktop"
OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tayyor")

# Fill these in. Anything left as ??? is printed as a warning at the end so
# no half-filled document goes to the centre by accident.
def _load_applicant():
    """Personal details, kept out of the repository.

    Name, home address and telephone number would otherwise sit in the git
    history of a public repository forever. They live in applicant.json,
    which is gitignored; applicant.example.json shows the shape.
    """
    import json
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "applicant.json")
    if not os.path.exists(path):
        raise SystemExit(
            "applicant.json topilmadi.\n"
            "  applicant.example.json dan nusxa oling va to'ldiring."
        )
    with open(path, encoding="utf-8") as handle:
        return json.load(handle)


APPLICANT = _load_applicant()

# ------------------------------------------------------- product content --
#
# Measured from the repository on 2026-08-21, not estimated:
#   40 982 lines of source (27 147 Dart + 13 835 Python)
#   106 server endpoints, 45 application screens, 17 data models
#   2,3 MB of source; ~338 MB with assets and models
PRODUCT = {
    "name_ru": ("SmartFlow — система автоматического учёта посещаемости и "
                "управления школой с распознаванием лиц и искусственным интеллектом"),
    "name_tg": ("Системаи худкори бақайдгирии ҳозиршавӣ ва идоракунии мактаб "
                "бо шиносоии чеҳра ва зеҳни сунъӣ «SmartFlow»"),
    "year": "2026 с.",
    "version": "Версия 1.0",
}

# Field code -> new value, for the big registration card (РК ИП).
FORM = {
    "3.1": PRODUCT["name_ru"],
    "3.2": PRODUCT["year"],
    "3.3": ("Платформа для автоматического учёта посещаемости учащихся по "
            "распознаванию лиц с камер, ведения электронного журнала оценок, "
            "выдачи онлайн-заданий и информирования родителей."),
    "3.4": ("Мобильное приложение (Android) + два серверных модуля "
            "(школьный и публичный) + база данных."),
    "3.5": PRODUCT["version"],
    "3.6": ("Образование, учёт посещаемости, компьютерное зрение, "
            "искусственный интеллект."),
    "3.7": ("Программный код (Dart/Flutter, Python/FastAPI), базы данных, "
            "модели машинного обучения для распознавания лиц."),
    "3.8": ("Мобильное приложение для четырёх ролей (директор, учитель, "
            "родитель, учащийся); школьный сервер с модулем распознавания "
            "лиц и электронным журналом; публичный сервер, через который "
            "родители и учащиеся получают доступ из дома."),
    "3.9": PRODUCT["year"],
    "3.10": "Мобильное приложение, серверные службы и база данных.",
    "3.11": "Образовательные учреждения.",
    "3.12": ("Школы, лицеи, гимназии, частные академии и учебные центры."),
    "3.13": ("Видеопоток IP-камер (распознавание лиц), ввод данных учителями "
             "и администрацией через мобильное приложение."),
    "3.14": "Ограниченный доступ (по регистрации и роли пользователя).",
    "3.15": "Таджикский, русский, английский.",
    "3.16": ("Записи посещаемости, оценки, учебные материалы, "
             "отчёты в формате PDF."),
    "3.17": ("338 МБ (исходный код — 2,3 МБ, 40 982 строки: "
             "27 147 строк Dart и 13 835 строк Python)."),
    "3.18": "Репозиторий исходного кода на платформе GitHub.",
    "3.19": ("Поиск реализован на стороне сервера средствами Python/SQLAlchemy "
             "по базе данных PostgreSQL. Поддерживается поиск и фильтрация по "
             "классам, учащимся, предметам, датам и статусу посещаемости без "
             "участия сторонних поисковых систем."),
    "3.20": ("Постоянно: данные посещаемости обновляются в режиме реального "
             "времени, функционал дорабатывается по мере эксплуатации."),
    "3.21": ("Локальный сервер школы (данные учащихся не выходят за пределы "
             "школьной сети) и публичный сервер для доступа родителей."),
    "3.23": ("Система состоит из трёх модулей.\n"
             "Школьный сервер — распознаёт лица учащихся с IP-камер и "
             "автоматически отмечает посещаемость; ведёт электронный журнал, "
             "расписание и учебные материалы. Фотографии и видеопоток "
             "учащихся не покидают этот сервер.\n"
             "Публичный сервер — получает только результат (оценки, "
             "посещаемость, объявления) и обслуживает родителей и учащихся "
             "через интернет.\n"
             "Мобильное приложение — единый интерфейс для четырёх ролей на "
             "трёх языках, с поддержкой тёмной темы."),

    "4.1": ("Процессор Intel Core i3 (4 ядра) или AMD Ryzen 3 и выше; "
            "оперативная память 8 ГБ; накопитель SSD 256 ГБ. "
            "Видеокарта не требуется — распознавание выполняется на "
            "центральном процессоре."),
    "4.2": ("Сервер: Windows 10/11 (64-бит). "
            "Приложение: Android 8.0 и выше."),
    "4.3": "PostgreSQL 14 и выше.",
    "4.5": ("Серверная часть — в локальной сети школы; "
            "мобильное приложение — установочный файл APK."),
    "4.6": ("Доступ разграничен по ролям: директор, учитель, родитель, "
            "учащийся. Каждый видит только те данные, которые относятся к нему."),

    # The basis is stated inside the cell rather than a bare figure. These
    # are the first installation's numbers, and a number whose derivation is
    # shown is one the centre can check; an unexplained one it can only take
    # on trust.
    "5.1": ("Более 200 пользователей: одна академия, 3 класса — учащиеся, "
            "их родители и преподаватели."),
    "5.2": ("Около 30 000 обращений в год (расчётно: ежедневная регистрация "
            "посещаемости учащихся в течение учебного года, а также обращения "
            "родителей и преподавателей через мобильное приложение)."),
    "5.3": ("Первое внедрение — частная академия: 3 класса, более 200 "
            "пользователей. Система поддерживает работу нескольких школ "
            "одновременно, при этом данные каждой школы полностью "
            "изолированы. Распознавание лиц выполняется локально, без "
            "передачи изображений в интернет."),

    "8.4": "Не проводилась",
}


def ownership_fields():
    """Sections 6 and 7 — owner and developer.

    These cells are rewritten whole rather than string-replaced. In the
    template each holds two people's details stacked in one cell, and the
    second person's e-mail is stored as a Word hyperlink, which a text
    replacement cannot see -- it survived every pass and would have gone to
    the centre as a joint owner's contact.
    """
    who = APPLICANT["fish"]
    contact = "Тел.: %s\nE-mail: %s" % (APPLICANT["phone"], APPLICANT["email"])
    coauthor = APPLICANT["coauthor"]

    if coauthor:
        who = "%s\n%s" % (who, coauthor)

    return {
        "6.1": who,
        "6.2": "-",
        "6.3": "Частное лицо" if not coauthor else "Частные лица",
        "6.4": "Физическое лицо — разработчик программного обеспечения",
        "6.5": APPLICANT["address_full"],
        "6.6": contact,
        "6.9": "Да" if coauthor else "Нет",
        "6.10": "Нет",
        "6.11": coauthor or APPLICANT["fish"],
        "7.1": who,
        "7.2": "-",
        "7.3": "Частное лицо" if not coauthor else "Частные лица",
        "7.4": APPLICANT["address_full"],
        "7.5": contact,
        "7.7": "-",
        "2.1": APPLICANT["fish"],
        "2.2": "-",
        "2.3": "Частное лицо",
        "2.4": APPLICANT["address"],
        "2.5": contact,
    }

# Everything that names the previous applicant or their product, replaced
# wherever it appears -- headings, tables and running text alike.
def replacements():
    fish = APPLICANT["fish"]
    # A second author only if there is one. Left over from the template, the
    # co-author's name would appear on the certificate as a joint owner.
    coauthor = APPLICANT["coauthor"] or ""
    return [
        # The centre is ММПИ (Маркази миллии патентӣ-иттилоотӣ). The
        # template names it НПИ throughout; addressing the application to
        # the wrong body is the kind of detail that gets papers returned.
        # Written as a stem so both "НПИ Центру" and "НПИ Центра РТ" match.
        ("НПИ Центр", "ММПИ Центр"),

        # The product name itself, which a blanket SmartDoc -> SmartFlow
        # swap left describing document management -- a different system.
        ("Система автоматизации документооборота школы c искусственным интеллектом SmartFlow",
         PRODUCT["name_ru"]),
        ("Система автоматизации документооборота школы c искусственным интеллектом SmartDoc",
         PRODUCT["name_ru"]),
        ("Системаи автоматикунонии ҳуҷҷатгузории мактаб бо зеҳни сунъӣ “SmartFlow”",
         "«%s»" % PRODUCT["name_tg"]),
        ("Системаи автоматикунонии ҳуҷҷатгузории мактаб бо зеҳни сунъӣ “SmartDoc”",
         "«%s»" % PRODUCT["name_tg"]),
        ("Мансурзода Ҳасан Мансур", coauthor),
        ("Мансурзода Хасан Мансур", coauthor),
        # The signature line uses initials, which the full-name rules miss.
        ("Мансурзода Х.М", APPLICANT["fish_short"] if not coauthor else coauthor),
        ("hachilov1986@gmail.com", coauthor and APPLICANT["email"] or ""),
        ("+992(92)840-11-15", coauthor and APPLICANT["phone"] or ""),
        ("735690, Таджикистан, Сугд, Бободжан Гафуров, Ёва, М. Рабиева 68", coauthor and APPLICANT["address_full"] or ""),
        ("Исполнительный комитет государственной власти Бободжон Гафуровского района", "-"),
        ("Рахимов Шероз Шарифжонович", fish),
        ("Раҳимов Шероз Шарифҷонович", fish),
        ("Рахимов Шероз", fish),
        ("Раҳимов Шероз", fish),
        ("Раҳимов Ш.Ш.", APPLICANT["fish_short"]),
        ("rahimovsheroz2009@gmail.com", APPLICANT["email"]),
        ("+992011775050", APPLICANT["phone"]),
        ("+992(01)177-50-50", APPLICANT["phone"]),
        ("735690, Республика Таджикистан, Бободжон Гафуровский, Хистеварз, ул. О.Хошимов б/н",
         APPLICANT["address_full"]),
        ("725690, Таджикистан, Сугд, Бободжан Гафуров, Хистеварз, О.Хошимов б/н",
         APPLICANT["address_full"]),
        ("Согдийская обл, Бободжан Гафуров, Хистеварз, О. Хошимов б/н",
         APPLICANT["address"]),
        ("SmartDoc", "SmartFlow"),
    ]


def replace_in_paragraph(paragraph, pairs):
    """Replaces across a paragraph without losing its formatting.

    Word splits text into runs at arbitrary points, so a phrase is often
    spread over several of them and a per-run replace misses it. The whole
    paragraph is rebuilt into its first run instead, which keeps that run's
    formatting for the line.
    """
    text = "".join(run.text for run in paragraph.runs)
    if not text:
        return False
    new = text
    for old, value in pairs:
        if old in new:
            new = new.replace(old, value)
    if new == text:
        return False
    paragraph.runs[0].text = new
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def set_cell(cell, value):
    """Writes a value into a table cell, keeping the first run's styling."""
    lines = value.split("\n")
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = lines[0]
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(lines[0])
    # Drop any paragraphs the template had beyond the first, then add ours.
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    for line in lines[1:]:
        cell.add_paragraph(line)


def process(filename, form_fields=None, rename=None):
    source = os.path.join(TEMPLATES, filename)
    target = os.path.join(OUTPUT, rename or filename.replace("Рахимов Шероз", "SmartFlow")
                                                    .replace("Рахимов_Шероз", "SmartFlow"))
    shutil.copyfile(source, target)

    document = docx.Document(target)
    pairs = replacements()
    changed = 0

    for paragraph in document.paragraphs:
        changed += replace_in_paragraph(paragraph, pairs)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changed += replace_in_paragraph(paragraph, pairs)

    # Then the field-by-field content, which must win over the blanket
    # replacements above.
    if form_fields:
        for table in document.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) < 3:
                    continue
                code = cells[0].text.strip()
                if code in form_fields:
                    set_cell(cells[-1], form_fields[code])
                    changed += 1

    document.save(target)
    print("  %-52s %d o'zgarish" % (os.path.basename(target), changed))
    return target


def main():
    os.makedirs(OUTPUT, exist_ok=True)
    print("Hujjatlar tayyorlanmoqda -> %s\n" % OUTPUT)

    # Only the two centre forms are template-filled. The description and the
    # two reviews are written from scratch by build_description.py and
    # build_reviews.py: their content is entirely different from the
    # previous applicant's, so copying those files would only carry another
    # product's text -- and another person's screenshots -- into this folder.
    process("Заявка_на_SmartDoc.docx", rename="Заявка_на_SmartFlow.docx")
    form = dict(FORM)
    form.update(ownership_fields())
    process("Ариза рус Рахимов Шероз.docx", form_fields=form)

    missing = [key for key, value in APPLICANT.items()
               if isinstance(value, str) and "???" in value]
    print()
    if missing:
        print("TO'LDIRILMAGAN MAYDONLAR (hujjatlarda ??? bo'lib qoladi):")
        for key in missing:
            print("   - %s" % key)
    else:
        print("Hamma maydon to'ldirilgan.")


if __name__ == "__main__":
    main()
